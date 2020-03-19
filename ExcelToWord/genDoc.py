# -*- coding: UTF-8 -*-
import os
import sys
import traceback

import xlrd
from docx import Document


def getCurrentPwd():
    return os.getcwd().replace('\\', '/').strip().rstrip('/')


current_pwd = getCurrentPwd()
source_file_name = getCurrentPwd() + '/数据表格.xlsx'
template_file_path = getCurrentPwd() + '/模板'
output_file_path = getCurrentPwd() + '/导出'


class DocxHandler():
    def __init__(self, filename, d_replace):
        self.document = Document(filename)
        self.d_replace = d_replace

    def replaceContent(self, content):
        old_content = content
        for k, v in self.d_replace.items():
            content = content.replace(k, v)
        return (old_content == content), content

    def replaceParagraphRuns(self, paragraph):
        if 0 == len(paragraph.runs): return
        run_text = ''
        for run in paragraph.runs:
            is_same, content = self.replaceContent(run.text)
            if not is_same: run.text = content
            run_text += run.text
        is_same, content = self.replaceContent(run_text)
        if not is_same:
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = content
                else:
                    run.text = ''

    def replaceParagraphText(self):
        for paragraph in self.document.paragraphs:
            print("---------替换段落内容---------")
            self.replaceParagraphRuns(paragraph)

    def replaceTableText(self):
        for table in self.document.tables:
            print("---------替换表格内容---------")
            for row in table.rows:
                for cell in row.cells:
                    if len(cell.paragraphs):
                        for paragraph in cell.paragraphs:
                            self.replaceParagraphRuns(paragraph)

    def replaceText(self):
        self.replaceTableText()
        self.replaceParagraphText()

    def save(self, filename=''):
        if filename:
            self.document.save(filename)


def doHandleDoc(wordApp, template_filename, output_filename, d_params):
    wordApp.openDocFile(template_filename)
    wordApp.replaceAllContent(d_params)
    wordApp.save(output_filename)
    wordApp.closeDocument()


def readSourceData():
    d_params = {}
    xls_fd = xlrd.open_workbook(source_file_name)
    sheet0 = xls_fd.sheet_by_index(0)
    row_count = sheet0.nrows
    for row in range(1, row_count):
        param = '${' + str(sheet0.cell(row, 1).value).strip() + '}'
        # param = str(sheet0.cell(row, 1).value).strip()
        value = str(sheet0.cell(row, 2).value).strip()
        d_params[param] = value
    return d_params


def entrance():
    d_params = readSourceData()
    template_file_list = []
    for root, dirs, files in os.walk(template_file_path):
        for file in files:
            file_type = os.path.splitext(file)[1]
            if file_type == '.docx':
                template_file_list.append(file)
    if len(template_file_list) == 0: return
    for file_name in template_file_list:
        if file_name.startswith('~$'): continue
        template_filename = template_file_path + '/' + file_name
        output_filename = output_file_path + '/' + file_name
        print('处理文件：', template_filename)
        doc = DocxHandler(template_filename, d_params)
        doc.replaceText()
        doc.save(output_filename)
        print('文件处理完成：', output_filename)
        # doHandleDoc(wordApp, template_filename, output_filename, d_params)
    # del wordApp


if __name__ == '__main__':
    args = len(sys.argv)
    file_name = ''
    if args >= 2:
        file_name = sys.argv[1]
    if file_name and not os.path.exists(file_name):
        file_name = getCurrentPwd() + '/' + file_name
        if not os.path.exists(file_name):
            file_name = source_file_name
            if not os.path.exists(file_name):
                print("数据文件不存在【%s】", file_name)
    if file_name: source_file_name = file_name
    if not os.path.exists(output_file_path): os.mkdir(output_file_path)

    try:
        entrance()
    except:
        print(traceback.format_exc())

    input()
